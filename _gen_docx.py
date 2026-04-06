from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── Helper: code block ───────────────────────────────────────────────────────
def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

# ── Helper: table with header row ────────────────────────────────────────────
def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D6E4F0')
        hdr_cells[i].paragraphs[0]._p.get_or_add_pPr().append(shading)
    # data rows
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
    return t

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('EDWH Framework — Validation Rules Change Manual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Date: 2026-04-01   |   Schema: GPC_DM   |   Author: EDWH Framework Team').italic = True

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Overview', 1)
doc.add_paragraph(
    'The C# backend (GPC Application) enforces data quality rules via FluentValidation validators '
    'whenever users upload Excel templates. The same rules have now been ported into the Oracle EDWH '
    'ETL pipeline so that data arriving from any source (including KBR_IHUB) is validated before '
    'being loaded into GPC_DM dimension tables.'
)

# ════════════════════════════════════════════════════════════════════════════
# 2. HOW VALIDATION WORKS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. How Validation Works in the Framework', 1)
doc.add_paragraph(
    'PKG_ETL_VALIDATOR.validate_staging reads rows from GPC_DM.ETL_VALIDATION_RULE for a given '
    'mapping and applies them to staging table rows with STG_STATUS = \'PENDING\'.'
)

doc.add_heading('Rule Types', 2)
add_table(doc,
    ['RULE_TYPE', 'Behaviour'],
    [
        ['NOT_NULL', 'Rejects staging rows where the column is NULL'],
        ['DERIVED',  'Attempts to populate the column via DERIVED_SQL; rejects if still NULL after derivation'],
        ['CHECK',    'Rejects staging rows where the SQL predicate stored in DERIVED_SQL evaluates to FALSE or NULL'],
    ]
)
doc.add_paragraph()

doc.add_heading('Error Actions', 2)
add_table(doc,
    ['ERROR_ACTION', 'Behaviour'],
    [
        ['REJECT', 'Marks the row REJECTED, appends reason to STG_REJECT_REASON, continues processing'],
        ['FAIL',   'Logs error, raises RAISE_APPLICATION_ERROR(-20032), aborts the entire run'],
    ]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. FILES CHANGED
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Files Changed', 1)

doc.add_heading('3.1  ddl/02_metadata_tables.sql', 2)
doc.add_paragraph(
    'The CHK_ETL_VR_TYPE constraint on ETL_VALIDATION_RULE.RULE_TYPE was extended to allow the value \'CHECK\'.'
)
add_code(doc, "-- BEFORE\nCONSTRAINT CHK_ETL_VR_TYPE CHECK (RULE_TYPE IN ('NOT_NULL','DERIVED','CUSTOM'))\n\n"
              "-- AFTER\nCONSTRAINT CHK_ETL_VR_TYPE CHECK (RULE_TYPE IN ('NOT_NULL','DERIVED','CUSTOM','CHECK'))")
doc.add_paragraph('The RULE_TYPE column comment was also updated to document the CHECK behaviour.')

doc.add_heading('3.2  packages/11_pkg_etl_validator.sql', 2)
doc.add_paragraph(
    'A new ELSIF branch was added inside the rule-type dispatch block in validate_staging to handle '
    'RULE_TYPE = \'CHECK\'. It executes a dynamic UPDATE rejecting PENDING rows where NOT (DERIVED_SQL) is true. '
    'On FAIL action, error code -20032 is raised.'
)
add_code(doc,
    "ELSIF vr.RULE_TYPE = 'CHECK' THEN\n"
    "    v_sql :=\n"
    "        'UPDATE ' || v_stg_table ||\n"
    "        ' SET STG_STATUS = ''REJECTED'',...'\n"
    "        ||' WHERE STG_RUN_ID = :run_id'\n"
    "        ||' AND   STG_STATUS = ''PENDING'''\n"
    "        ||' AND   NOT (' || vr.DERIVED_SQL || ')';\n"
    "    EXECUTE IMMEDIATE v_sql USING vr.RULE_NAME, p_run_id;"
)

doc.add_heading('3.3  data/16_metadata_inserts.sql', 2)
doc.add_paragraph('14 new INSERT INTO GPC_DM.ETL_VALIDATION_RULE statements were added across four mapping sections.')

# ════════════════════════════════════════════════════════════════════════════
# 4. NEW VALIDATION RULES
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. New Validation Rules', 1)

rule_sections = [
    (
        '4.1  STAFFING_SCHEDULE → DIM_STAFFING_SCHEDULE',
        'Mirrors: TimelineExelValidator (CONTRACT_TYPE, STATUS, date order) and PaffExelValidator (PAF_STATUS, CONTRACT_TYPE)',
        [
            ('PROJECT_ID_REQUIRED',  'NOT_NULL', 'PROJECT_ID',         'Must not be NULL',                                          'REJECT'),
            ('POSITION_ID_REQUIRED', 'NOT_NULL', 'POSITION_ID',        'Must not be NULL',                                          'REJECT'),
            ('SCHEDULE_TYPE_CHECK',  'CHECK',    'SCHEDULE_TYPE',      "NULL or IN ('Exempt','Non-Exempt','Exempt Agency','Non-Exempt Agency')", 'REJECT'),
            ('SCHEDULE_STATUS_CHECK','CHECK',    'SCHEDULE_STATUS',    "NULL or IN ('Open','Filled','Canceled','Approved','Rejected','Withdrawn','Pending')", 'REJECT'),
            ('SCHEDULE_DATE_ORDER',  'CHECK',    'SCHEDULE_START_DATE','NULL or START_DATE <= END_DATE',                             'REJECT'),
        ]
    ),
    (
        '4.2  STAFFING_SCHEDULE → DIM_STAFFING_TIMELINE',
        'Mirrors: TimelineExelValidator (HOURS_PER_WEEK numeric, date order)',
        [
            ('ALLOCATED_HOURS_REQUIRED',     'NOT_NULL', 'ALLOCATED_HOURS',   'Must not be NULL',                          'REJECT'),
            ('ALLOCATED_HOURS_NON_NEGATIVE', 'CHECK',    'ALLOCATED_HOURS',   'NULL or >= 0',                              'REJECT'),
            ('PERIOD_DATE_ORDER',            'CHECK',    'PERIOD_START_DATE', 'NULL or PERIOD_START_DATE <= PERIOD_END_DATE','REJECT'),
        ]
    ),
    (
        '4.3  COST → DIM_COST',
        'Mirrors: CostExcelValidator (COST_BASIS must be COMPANY or CLIENT)',
        [
            ('PROJECT_ID_REQUIRED', 'NOT_NULL', 'PROJECT_ID',    'Must not be NULL',                                 'REJECT'),
            ('COST_TYPE_REQUIRED',  'NOT_NULL', 'COST_TYPE',     'Must not be NULL',                                 'REJECT'),
            ('COST_CATEGORY_CHECK', 'CHECK',    'COST_CATEGORY', "NULL or UPPER(COST_CATEGORY) IN ('COMPANY','CLIENT')", 'REJECT'),
        ]
    ),
    (
        '4.4  COST → DIM_TIMELINE_COST',
        "Mirrors: CostExcelValidator (CURVE_TYPE enum). Note: EDWH classifies CURVE_TYPE as 'ACTUAL'/'BUDGET'/'FORECAST', not the scheduling distribution curve names used in the C# UI.",
        [
            ('AMOUNT_REQUIRED',        'NOT_NULL', 'AMOUNT',        'Must not be NULL',                                      'REJECT'),
            ('FORECAST_TYPE_REQUIRED', 'NOT_NULL', 'FORECAST_TYPE', 'Must not be NULL',                                      'REJECT'),
            ('CURVE_TYPE_VALUE_CHECK', 'CHECK',    'CURVE_TYPE',    "NULL or IN ('ACTUAL','BUDGET','FORECAST')",              'REJECT'),
        ]
    ),
]

for heading, note, rules in rule_sections:
    doc.add_heading(heading, 2)
    add_table(doc,
        ['Rule Name', 'Type', 'Column', 'Condition', 'Action'],
        [list(r) for r in rules]
    )
    p = doc.add_paragraph(note)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9.5)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. HOW TO APPLY
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. How to Apply to an Existing Oracle Environment', 1)

doc.add_heading('Scenario A — Fresh Install (tables do not yet exist)', 2)
doc.add_paragraph('Run the full install script in order:')
add_code(doc, '@00_install_all.sql')

doc.add_heading('Scenario B — Existing Install (tables already deployed)', 2)

doc.add_heading('Step 1 — Alter the RULE_TYPE constraint', 3)
add_code(doc,
    "ALTER TABLE GPC_DM.ETL_VALIDATION_RULE\n"
    "    DROP CONSTRAINT CHK_ETL_VR_TYPE;\n\n"
    "ALTER TABLE GPC_DM.ETL_VALIDATION_RULE\n"
    "    ADD CONSTRAINT CHK_ETL_VR_TYPE\n"
    "        CHECK (RULE_TYPE IN ('NOT_NULL','DERIVED','CUSTOM','CHECK'));"
)

doc.add_heading('Step 2 — Recompile the validator package', 3)
doc.add_paragraph('Run the updated package spec and body:')
add_code(doc, '@packages/11_pkg_etl_validator.sql')
doc.add_paragraph('Verify compilation:')
add_code(doc,
    "SELECT OBJECT_NAME, STATUS\n"
    "FROM   ALL_OBJECTS\n"
    "WHERE  OWNER       = 'GPC_DM'\n"
    "AND    OBJECT_NAME = 'PKG_ETL_VALIDATOR';\n"
    "-- Expected: STATUS = VALID"
)

doc.add_heading('Step 3 — Insert the new validation rules', 3)
doc.add_paragraph(
    'Run the metadata inserts file, or insert rules selectively. '
    'Retrieve current mapping IDs first:'
)
add_code(doc,
    "SELECT m.MAPPING_ID, e.ENTITY_NAME, m.TARGET_TABLE\n"
    "FROM   GPC_DM.ETL_TARGET_MAPPING m\n"
    "JOIN   GPC_DM.ETL_ENTITY         e ON e.ENTITY_ID = m.ENTITY_ID\n"
    "ORDER BY m.MAPPING_ID;"
)
doc.add_paragraph('Then insert the 14 new rules using the actual MAPPING_ID values returned.')

doc.add_heading('Step 4 — Verify', 3)
add_code(doc,
    "SELECT RULE_ID, MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, ERROR_ACTION\n"
    "FROM   GPC_DM.ETL_VALIDATION_RULE\n"
    "WHERE  RULE_TYPE = 'CHECK'\n"
    "ORDER BY MAPPING_ID, RULE_ID;\n"
    "-- Expected: 8 CHECK rules"
)

# ════════════════════════════════════════════════════════════════════════════
# 6. RUNTIME BEHAVIOUR
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Validation Behaviour at Runtime', 1)
doc.add_paragraph(
    'When PKG_ETL_RUNNER processes a mapping it calls PKG_ETL_VALIDATOR.validate_staging. '
    'The function loops through all active rules in RULE_ID order. For each CHECK rule it executes:'
)
add_code(doc,
    "UPDATE <staging_table>\n"
    "SET    STG_STATUS        = 'REJECTED',\n"
    "       STG_REJECT_REASON = SUBSTR(NVL(STG_REJECT_REASON,'')\n"
    "                           || ' [<rule_name>: <column> check failed]', 1, 500)\n"
    "WHERE  STG_RUN_ID = <run_id>\n"
    "AND    STG_STATUS = 'PENDING'\n"
    "AND    NOT (<DERIVED_SQL predicate>)"
)
doc.add_paragraph('Rejected rows can be reviewed with:')
add_code(doc,
    "SELECT *\n"
    "FROM   GPC_DM.<staging_table>\n"
    "WHERE  STG_RUN_ID = <run_id>\n"
    "AND    STG_STATUS = 'REJECTED'\n"
    "ORDER BY STG_REJECT_REASON;"
)

# ════════════════════════════════════════════════════════════════════════════
# 7. C# ↔ SQL MAPPING TABLE
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Correspondence: C# Validators → SQL Rules', 1)
add_table(doc,
    ['C# Validator', 'C# Rule', 'SQL Rule Name', 'Target Mapping'],
    [
        ['TimelineExelValidator', 'CONTRACT_TYPE IN (Exempt, Non-Exempt, ...)', 'SCHEDULE_TYPE_CHECK',  'DIM_STAFFING_SCHEDULE'],
        ['TimelineExelValidator', 'STATUS IN (Open, Filled, Canceled)',          'SCHEDULE_STATUS_CHECK','DIM_STAFFING_SCHEDULE'],
        ['TimelineExelValidator', 'PlanStartDate <= PlanEndDate',                'SCHEDULE_DATE_ORDER',  'DIM_STAFFING_SCHEDULE'],
        ['TimelineExelValidator', 'HoursPerWeek numeric',                        'ALLOCATED_HOURS_NON_NEGATIVE', 'DIM_STAFFING_TIMELINE'],
        ['PaffExelValidator',     'CONTRACT_TYPE enum',                          'SCHEDULE_TYPE_CHECK',  'DIM_STAFFING_SCHEDULE'],
        ['PaffExelValidator',     'PAF_STATUS enum',                             'SCHEDULE_STATUS_CHECK','DIM_STAFFING_SCHEDULE'],
        ['PaffExelValidator',     'StartDate <= EndDate',                        'SCHEDULE_DATE_ORDER',  'DIM_STAFFING_SCHEDULE'],
        ['CostExcelValidator',    'COST_BASIS IN (COMPANY, CLIENT)',             'COST_CATEGORY_CHECK',  'DIM_COST'],
        ['CostExcelValidator',    'CURVE_TYPE enum',                             'CURVE_TYPE_VALUE_CHECK','DIM_TIMELINE_COST'],
    ]
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 8. NOTES
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Notes', 1)
notes = [
    'All current rules use ERROR_ACTION = REJECT, allowing the ETL run to continue. Only invalid rows are excluded from the load.',
    'To abort a run on any violation, update ERROR_ACTION = \'FAIL\' in ETL_VALIDATION_RULE for the relevant rule.',
    'New rules can be added at any time by inserting rows into ETL_VALIDATION_RULE — no package recompilation required.',
    'CHECK predicates in DERIVED_SQL must be valid Oracle SQL WHERE-clause expressions over the staging table columns.',
    'NULL handling: optional fields use COLUMN IS NULL OR COLUMN IN (...) so NULL values are not rejected by CHECK (use NOT_NULL rules to enforce required fields separately).',
]
for note in notes:
    p = doc.add_paragraph(note, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(r'c:\Aralytiks\AR-Dev\EDWH_Framework\VALIDATION_CHANGES_MANUAL.docx')
print("Done: VALIDATION_CHANGES_MANUAL.docx created.")
