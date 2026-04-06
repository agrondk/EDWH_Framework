"""
Generates SOLUTION_MANUAL.docx from the Markdown content.
Run: python _gen_solution_manual_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE   = RGBColor(0x1F, 0x49, 0x7D)
DKBLUE = RGBColor(0x0D, 0x2B, 0x55)
GREY   = RGBColor(0x44, 0x44, 0x44)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
HDR_BG = 'D6E4F0'
CODE_BG = 'F4F4F4'

doc = Document()

# ─── Page margins ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Base font ──────────────────────────────────────────────────────────────
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

for lvl, sz in [(1, 16), (2, 13), (3, 11)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Calibri'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = BLUE

# ─── Helpers ────────────────────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def shade_para(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def add_code(text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = GREY
        shade_para(p, CODE_BG)
    doc.add_paragraph()  # spacer

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False

    # Set column widths
    usable = 16.0  # cm
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    else:
        each = usable / len(headers)
        for col in t.columns:
            for cell in col.cells:
                cell.width = Cm(each)

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, 'D6E4F0')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DKBLUE

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r + 1].cells[c]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            if r % 2 == 1:
                shade_cell(cell, 'F7FAFD')

    doc.add_paragraph()

def h1(text):  doc.add_heading(text, 1)
def h2(text):  doc.add_heading(text, 2)
def h3(text):  doc.add_heading(text, 3)
def para(text, italic=False, bold=False, sz=None):
    p = doc.add_paragraph(text)
    if italic or bold or sz:
        for run in p.runs:
            run.italic = italic
            run.bold   = bold
            if sz: run.font.size = Pt(sz)
    return p
def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_heading('GPC_DM EDWH Framework', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = DKBLUE

sub = doc.add_heading('Solution Manual', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = BLUE

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Version 1.0   |   Date: 2026-04-01   |   Schema: GPC_DM   |   Source: KBR_IHUB').italic = True
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 1. SOLUTION OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
h1('1. Solution Overview')
para(
    'The GPC_DM EDWH Framework is a metadata-driven Oracle ETL pipeline that loads data from '
    'upstream source systems (currently KBR_IHUB) into the GPC_DM data warehouse dimension tables.'
)
for b in [
    'Metadata-driven — source tables, column mappings, load types, and validation rules are stored in registry tables. No code changes needed to add columns or adjust mappings.',
    'Validation parity — the same data quality rules enforced by the C# backend (FluentValidation) on Excel uploads are also enforced by the SQL pipeline from any source.',
    'Idempotent / restartable — failed runs can be safely re-run. Staging tables are cleared and re-populated at the start of each run.',
    'Full audit trail — every run, step, row count, rejection reason, and error is written to logging tables that survive rollback (autonomous transactions).',
    'SCD2 + Incremental — slowly changing dimension logic for master records; incremental MERGE for timeline/period data.',
]:
    bullet(b)

doc.add_paragraph()
h2('Current Entities')
add_table(
    ['Entity Name', 'Source Table', 'Target Tables', 'Load Type'],
    [
        ['STAFFING_SCHEDULE', 'KBR_IHUB.APAC_PCDM_STAFFING_SCHEDULE_IHUB', 'DIM_STAFFING_SCHEDULE', 'SCD2'],
        ['',                  '',                                             'DIM_STAFFING_TIMELINE',  'INCREMENTAL'],
        ['COST',              'KBR_IHUB.APAC_PCDM_COST_IHUB',               'DIM_COST',               'SCD2'],
        ['',                  '',                                             'DIM_TIMELINE_COST',      'INCREMENTAL'],
    ],
    col_widths=[3.8, 5.8, 4.0, 2.4]
)


# ════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
h1('2. Architecture')
h2('ETL Flow')
add_code(
    "KBR_IHUB source tables\n"
    "        │\n"
    "        ▼\n"
    " PKG_ETL_ORCHESTRATOR.run_entity()\n"
    "        │\n"
    "        ├─ 1. Validate metadata\n"
    "        ├─ 2. Acquire lock (ETL_CONTROL)\n"
    "        ├─ 3. Determine watermark window\n"
    "        │\n"
    "        ├─ 4. TRANSFORM (entity-specific package)\n"
    "        │       └─ Extract → deduplicate → hash → INSERT into staging tables\n"
    "        │\n"
    "        └─ 5. For each target mapping:\n"
    "               ├─ a. VALIDATE STAGING (PKG_ETL_VALIDATOR)\n"
    "               │       └─ NOT_NULL / DERIVED / CHECK rules\n"
    "               ├─ b. LOAD — SCD2: CLASSIFY→EXPIRE→INSERT\n"
    "               │         — INCREMENTAL: MERGE on business key\n"
    "               ├─ c. Post-load SCD2 duplicate check\n"
    "               └─ d. COMMIT + advance watermark"
)

h2('Staging Table Status Values')
add_table(
    ['STG_STATUS', 'Meaning'],
    [
        ['PENDING',  'Loaded by transform; awaiting validation and load'],
        ['LOADED',   'Successfully written to the target DIM table'],
        ['REJECTED', 'Failed validation; not loaded; reason in STG_REJECT_REASON'],
    ],
    col_widths=[3, 13]
)

h2('ETL Control States')
add_table(
    ['Status', 'Meaning', 'Transition'],
    [
        ['IDLE',     'Ready to run',                     'Set to RUNNING at run start'],
        ['RUNNING',  'Run in progress (locked)',          'Returns to IDLE (success) or FAILED (error)'],
        ['FAILED',   'Last run failed',                   'Retried by run_all; or manually reset to IDLE'],
        ['DISABLED', 'Excluded from run_all',             'Manually set; re-enable by setting to IDLE'],
    ],
    col_widths=[2.5, 5.5, 8.0]
)

h2('SCD2 Control Columns (DIM tables)')
add_table(
    ['Column', 'Meaning'],
    [
        ['EFFECTIVE_START_DATE', 'Date this row version became active'],
        ['EFFECTIVE_END_DATE',   '9999-12-31 = currently active; otherwise = superseded date'],
        ['IS_CURRENT',           'Y = active version; N = historical'],
        ['RECORD_HASH',          'SHA-256 hash of tracked columns; used for change detection'],
    ],
    col_widths=[4.5, 11.5]
)


# ════════════════════════════════════════════════════════════════════════════
# 3. FILE INVENTORY
# ════════════════════════════════════════════════════════════════════════════
h1('3. File Inventory')
add_table(
    ['File', 'Purpose'],
    [
        ['00_install_all.sql',                    'Master install script — runs all files in order'],
        ['ddl/01_sequences.sql',                  'All sequences for surrogate keys and log IDs'],
        ['ddl/02_metadata_tables.sql',            'Registry tables: ETL_SOURCE_SYSTEM, ETL_ENTITY, ETL_TARGET_MAPPING, ETL_COLUMN_MAPPING, ETL_VALIDATION_RULE'],
        ['ddl/03_control_tables.sql',             'ETL_CONTROL — watermark and run state per entity'],
        ['ddl/04_logging_tables.sql',             'ETL_RUN_LOG, ETL_STEP_LOG, ETL_ERROR_LOG'],
        ['ddl/05_target_tables.sql',              'DIM_STAFFING_SCHEDULE, DIM_STAFFING_TIMELINE, DIM_COST, DIM_TIMELINE_COST'],
        ['ddl/06_staging_tables.sql',             'STG_STAFFING_SCHEDULE, STG_STAFFING_TIMELINE, STG_COST, STG_TIMELINE_COST'],
        ['ddl/07_views.sql',                      'V_ETL_ACTIVE_ENTITIES, V_ETL_RUN_SUMMARY'],
        ['packages/08_pkg_etl_logger.sql',        'Logging package; all writes via autonomous transactions'],
        ['packages/09_pkg_etl_metadata.sql',      'Metadata read helpers (get_entity, get_mappings, get_bk_cols)'],
        ['packages/10_pkg_etl_control.sql',       'Watermark, lock, entity state management'],
        ['packages/11_pkg_etl_validator.sql',     'Staging validation: NOT_NULL, DERIVED, CHECK rules'],
        ['packages/12_pkg_etl_scd2_loader.sql',   'Generic SCD2 loader and incremental MERGE loader'],
        ['packages/13_pkg_etl_transform_staffing.sql', 'Transform package for STAFFING_SCHEDULE entity'],
        ['packages/14_pkg_etl_transform_cost.sql','Transform package for COST entity'],
        ['packages/15_pkg_etl_orchestrator.sql',  'Main entry point: run_entity(), run_all(), dispatch_transform()'],
        ['data/16_metadata_inserts.sql',          'Seed data: entities, mappings, column mappings, validation rules'],
        ['monitoring/17_monitoring_queries.sql',  'Operational monitoring SQL queries'],
    ],
    col_widths=[5.5, 10.5]
)


# ════════════════════════════════════════════════════════════════════════════
# 4. PREREQUISITES
# ════════════════════════════════════════════════════════════════════════════
h1('4. Prerequisites')
add_table(
    ['Requirement', 'Details'],
    [
        ['Oracle version',    '12c or later (uses GENERATED ALWAYS AS IDENTITY, STANDARD_HASH)'],
        ['Target schema',     'GPC_DM must exist'],
        ['Source schema',     'KBR_IHUB must be accessible from the GPC_DM connection'],
        ['Privileges',        'CREATE TABLE, CREATE SEQUENCE, CREATE VIEW, CREATE PROCEDURE on GPC_DM'],
        ['Source access',     'SELECT on KBR_IHUB.APAC_PCDM_STAFFING_SCHEDULE_IHUB and KBR_IHUB.APAC_PCDM_COST_IHUB granted to GPC_DM'],
        ['SQL client',        'SQL*Plus or SQLcl (scripts use @@ for relative includes)'],
    ],
    col_widths=[3.5, 12.5]
)
para('Verify source access before installing:')
add_code(
    "SELECT COUNT(*) FROM KBR_IHUB.APAC_PCDM_STAFFING_SCHEDULE_IHUB;\n"
    "SELECT COUNT(*) FROM KBR_IHUB.APAC_PCDM_COST_IHUB;\n"
    "-- Both must return without error"
)


# ════════════════════════════════════════════════════════════════════════════
# 5. INSTALLATION — FRESH ENVIRONMENT
# ════════════════════════════════════════════════════════════════════════════
h1('5. Installation — Fresh Environment')

h2('Step 1 — Connect as GPC_DM owner')
add_code('CONNECT gpc_dm/<password>@<tns_alias>')

h2('Step 2 — Navigate to the framework folder')
para('The @@ calls in 00_install_all.sql use relative paths. Run the script from the EDWH_Framework directory.')
add_code('-- SQLcl\ncd /path/to/EDWH_Framework')

h2('Step 3 — Run the master install script')
add_code('@00_install_all.sql')
para('The script uses WHENEVER SQLERROR EXIT FAILURE ROLLBACK — any failure stops immediately. Fix the error and re-run.')

h2('Step 4 — Verify Installation')
add_code(
    "-- Check all packages compiled\n"
    "SELECT OBJECT_NAME, STATUS FROM ALL_OBJECTS\n"
    "WHERE  OWNER = 'GPC_DM' AND OBJECT_TYPE IN ('PACKAGE','PACKAGE BODY')\n"
    "ORDER BY OBJECT_NAME;\n"
    "-- All rows must show STATUS = VALID\n\n"
    "-- Check entities registered\n"
    "SELECT * FROM GPC_DM.V_ETL_ACTIVE_ENTITIES;\n\n"
    "-- Check validation rules (expect 14 rows)\n"
    "SELECT RULE_NAME, RULE_TYPE, COLUMN_NAME FROM GPC_DM.ETL_VALIDATION_RULE ORDER BY MAPPING_ID, RULE_ID;"
)


# ════════════════════════════════════════════════════════════════════════════
# 6. INSTALLATION — INCREMENTAL
# ════════════════════════════════════════════════════════════════════════════
h1('6. Installation — Existing Environment (Incremental Apply)')

h2('6.1 — Alter RULE_TYPE Constraint')
para('Only needed if the environment was installed before the CHECK rule type was added.')
add_code(
    "ALTER TABLE GPC_DM.ETL_VALIDATION_RULE DROP CONSTRAINT CHK_ETL_VR_TYPE;\n\n"
    "ALTER TABLE GPC_DM.ETL_VALIDATION_RULE\n"
    "    ADD CONSTRAINT CHK_ETL_VR_TYPE\n"
    "        CHECK (RULE_TYPE IN ('NOT_NULL','DERIVED','CUSTOM','CHECK'));"
)

h2('6.2 — Recompile Updated Packages')
add_code(
    "@packages/11_pkg_etl_validator.sql\n"
    "@packages/15_pkg_etl_orchestrator.sql\n\n"
    "-- Verify\n"
    "SELECT OBJECT_NAME, STATUS FROM ALL_OBJECTS\n"
    "WHERE  OWNER = 'GPC_DM' AND OBJECT_TYPE = 'PACKAGE BODY' AND STATUS != 'VALID';\n"
    "-- Must return no rows"
)

h2('6.3 — Apply New Metadata')
add_code(
    "-- Find existing mapping IDs\n"
    "SELECT m.MAPPING_ID, e.ENTITY_NAME, m.TARGET_TABLE\n"
    "FROM   GPC_DM.ETL_TARGET_MAPPING m\n"
    "JOIN   GPC_DM.ETL_ENTITY         e ON e.ENTITY_ID = m.ENTITY_ID\n"
    "ORDER BY m.MAPPING_ID;\n\n"
    "-- Then insert new rules using actual MAPPING_IDs\n"
    "INSERT INTO GPC_DM.ETL_VALIDATION_RULE (MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, ERROR_ACTION)\n"
    "VALUES (<mapping_id>, 'MY_NEW_RULE', 'NOT_NULL', 'MY_COLUMN', 'REJECT');\n"
    "COMMIT;"
)

h2('6.4 — Add New Columns')
add_code(
    "ALTER TABLE GPC_DM.DIM_STAFFING_SCHEDULE ADD (NEW_COL VARCHAR2(100));\n"
    "ALTER TABLE GPC_DM.STG_STAFFING_SCHEDULE ADD (NEW_COL VARCHAR2(100));\n\n"
    "INSERT INTO GPC_DM.ETL_COLUMN_MAPPING\n"
    "    (MAPPING_ID, SOURCE_COLUMN, TARGET_COLUMN, IS_BUSINESS_KEY, IS_TRACKED, COLUMN_ORDER)\n"
    "VALUES (<mapping_id>, 'SOURCE_NEW_COL', 'NEW_COL', 'N', 'Y', <next_order>);\n"
    "COMMIT;"
)


# ════════════════════════════════════════════════════════════════════════════
# 7. RUNNING THE ETL
# ════════════════════════════════════════════════════════════════════════════
h1('7. Running the ETL')

h2('Run a Single Entity')
add_code(
    "EXEC GPC_DM.PKG_ETL_ORCHESTRATOR.run_entity('STAFFING_SCHEDULE');\n"
    "EXEC GPC_DM.PKG_ETL_ORCHESTRATOR.run_entity('COST');"
)

h2('Run All Active Entities')
add_code("EXEC GPC_DM.PKG_ETL_ORCHESTRATOR.run_all;")
para('Processes entities in ENTITY_ID order. A failure in one entity is logged; others continue.')

h2('Schedule with Oracle DBMS_SCHEDULER')
add_code(
    "BEGIN\n"
    "    DBMS_SCHEDULER.CREATE_JOB(\n"
    "        job_name        => 'GPC_DM.ETL_DAILY_RUN',\n"
    "        job_type        => 'PLSQL_BLOCK',\n"
    "        job_action      => 'BEGIN GPC_DM.PKG_ETL_ORCHESTRATOR.run_all; END;',\n"
    "        start_date      => SYSTIMESTAMP,\n"
    "        repeat_interval => 'FREQ=DAILY;BYHOUR=2;BYMINUTE=0',\n"
    "        enabled         => TRUE\n"
    "    );\n"
    "END;\n/"
)

h2('Watermark Override (Reprocess Historical Data)')
add_code(
    "-- Reprocess from a specific date\n"
    "UPDATE GPC_DM.ETL_CONTROL\n"
    "SET    LAST_WATERMARK = DATE '2025-01-01'\n"
    "WHERE  ENTITY_ID = (SELECT ENTITY_ID FROM GPC_DM.ETL_ENTITY WHERE ENTITY_NAME = 'COST');\n"
    "COMMIT;\n\n"
    "-- Full reload: set watermark to NULL\n"
    "UPDATE GPC_DM.ETL_CONTROL SET LAST_WATERMARK = NULL WHERE ENTITY_ID = ...;\n"
    "COMMIT;"
)


# ════════════════════════════════════════════════════════════════════════════
# 8. MONITORING
# ════════════════════════════════════════════════════════════════════════════
h1('8. Monitoring and Operations')

h2('Entity Status')
add_code(
    "SELECT e.ENTITY_NAME, c.STATUS, c.LAST_WATERMARK, c.LAST_RUN_DATE\n"
    "FROM   GPC_DM.ETL_CONTROL c JOIN GPC_DM.ETL_ENTITY e ON e.ENTITY_ID = c.ENTITY_ID\n"
    "ORDER BY e.ENTITY_NAME;"
)

h2('Recent Runs')
add_code("SELECT * FROM GPC_DM.V_ETL_RUN_SUMMARY WHERE ROWNUM <= 10;")

h2('Errors in Last 24 Hours')
add_code(
    "SELECT RUN_ID, ENTITY_NAME, ERROR_CODE, ERROR_MESSAGE, ERROR_TIME\n"
    "FROM   GPC_DM.ETL_ERROR_LOG\n"
    "WHERE  ERROR_TIME >= SYSDATE - 1\n"
    "ORDER BY ERROR_TIME DESC;"
)

h2('Step Detail for a Specific Run')
add_code(
    "SELECT STEP_NAME, STATUS, ROWS_AFFECTED, START_TIME, STEP_MESSAGE\n"
    "FROM   GPC_DM.ETL_STEP_LOG\n"
    "WHERE  RUN_ID = <your_run_id>\n"
    "ORDER BY STEP_ID;"
)

h2('Rejected Rows')
add_code(
    "SELECT STG_ID, SCHEDULE_ID, STG_REJECT_REASON\n"
    "FROM   GPC_DM.STG_STAFFING_SCHEDULE\n"
    "WHERE  STG_RUN_ID = <run_id> AND STG_STATUS = 'REJECTED';"
)


# ════════════════════════════════════════════════════════════════════════════
# 9. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
h1('9. Troubleshooting and Recovery')

h2('Entity Stuck in RUNNING State')
para('If a session was killed mid-run, the entity stays RUNNING and cannot be re-run.')
add_code(
    "UPDATE GPC_DM.ETL_CONTROL SET STATUS = 'IDLE'\n"
    "WHERE  ENTITY_ID = (SELECT ENTITY_ID FROM GPC_DM.ETL_ENTITY WHERE ENTITY_NAME = 'STAFFING_SCHEDULE');\n"
    "COMMIT;"
)

h2('Package Compilation Errors')
add_code(
    "SELECT TEXT FROM ALL_ERRORS\n"
    "WHERE  OWNER = 'GPC_DM' AND NAME = 'PKG_ETL_VALIDATOR'\n"
    "ORDER BY SEQUENCE;"
)

h2('Temporarily Disable a Validation Rule')
add_code(
    "UPDATE GPC_DM.ETL_VALIDATION_RULE SET IS_ACTIVE = 'N' WHERE RULE_NAME = 'SCHEDULE_TYPE_CHECK';\n"
    "COMMIT;\n"
    "-- Re-enable:\n"
    "UPDATE GPC_DM.ETL_VALIDATION_RULE SET IS_ACTIVE = 'Y' WHERE RULE_NAME = 'SCHEDULE_TYPE_CHECK';\n"
    "COMMIT;"
)

h2('Disable an Entity from run_all')
add_code(
    "UPDATE GPC_DM.ETL_CONTROL\n"
    "SET    STATUS = 'DISABLED', NOTES = 'Source migration in progress'\n"
    "WHERE  ENTITY_ID = (SELECT ENTITY_ID FROM GPC_DM.ETL_ENTITY WHERE ENTITY_NAME = 'COST');\n"
    "COMMIT;"
)


# ════════════════════════════════════════════════════════════════════════════
# 10. ADD A NEW SOURCE SYSTEM
# ════════════════════════════════════════════════════════════════════════════
h1('10. How to Add a New Source System')

h2('Step 1 — Register the Source System')
add_code(
    "INSERT INTO GPC_DM.ETL_SOURCE_SYSTEM (SS_NAME, SS_SCHEMA, SS_DESCRIPTION)\n"
    "VALUES ('MY_SOURCE', 'MY_SOURCE_SCHEMA', 'Description of the new source system');\n"
    "COMMIT;"
)

h2('Step 2 — Grant SELECT Privileges')
add_code("GRANT SELECT ON MY_SOURCE_SCHEMA.MY_TABLE TO GPC_DM;")

para('Then follow Section 11 to register entities from this source system.')


# ════════════════════════════════════════════════════════════════════════════
# 11. ADD A NEW ENTITY / TABLE
# ════════════════════════════════════════════════════════════════════════════
h1('11. How to Add a New Entity / Table')

para('Adding a new entity requires four areas of change: DDL, transform package, metadata registration, and orchestrator dispatch.')

steps = [
    ('1', 'Create DDL', 'Target DIM table, staging STG table, surrogate key sequence'),
    ('2', 'Transform package', 'New PKG_ETL_TRANSFORM_<ENTITY>.transform() procedure'),
    ('3', 'Register metadata', 'ETL_ENTITY, ETL_CONTROL, ETL_TARGET_MAPPING, ETL_COLUMN_MAPPING, ETL_VALIDATION_RULE'),
    ('4', 'Orchestrator dispatch', 'Add WHEN branch in PKG_ETL_ORCHESTRATOR.dispatch_transform()'),
]
add_table(['Step', 'What', 'Details'], steps, col_widths=[1.0, 3.5, 11.5])

h2('Step 1 — Create DDL')
para('Sequence:')
add_code("CREATE SEQUENCE GPC_DM.SEQ_DIM_MY_ENTITY START WITH 1 INCREMENT BY 1 NOCACHE NOCYCLE;")

para('Target table (SCD2 example):')
add_code(
    "CREATE TABLE GPC_DM.DIM_MY_ENTITY (\n"
    "    DIM_ME_ID            NUMBER         NOT NULL,\n"
    "    MY_ENTITY_ID         VARCHAR2(50)   NOT NULL,  -- Business key\n"
    "    ATTRIBUTE_1          VARCHAR2(100),\n"
    "    ATTRIBUTE_2          DATE,\n"
    "    EFFECTIVE_START_DATE DATE           NOT NULL,\n"
    "    EFFECTIVE_END_DATE   DATE           DEFAULT DATE '9999-12-31' NOT NULL,\n"
    "    IS_CURRENT           VARCHAR2(1)    DEFAULT 'Y' NOT NULL,\n"
    "    RECORD_HASH          VARCHAR2(64),\n"
    "    REPORTING_DATE       DATE,\n"
    "    ETL_RUN_ID           NUMBER,\n"
    "    ETL_LOAD_DATE        DATE           DEFAULT SYSDATE,\n"
    "    CONSTRAINT PK_DIM_ME          PRIMARY KEY (DIM_ME_ID),\n"
    "    CONSTRAINT UQ_DIM_ME_EFFSTART UNIQUE (MY_ENTITY_ID, EFFECTIVE_START_DATE)\n"
    ");\n"
    "CREATE INDEX GPC_DM.IDX_DIM_ME_BK ON GPC_DM.DIM_MY_ENTITY(MY_ENTITY_ID, IS_CURRENT);"
)

para('Staging table:')
add_code(
    "CREATE TABLE GPC_DM.STG_MY_ENTITY (\n"
    "    STG_ID            NUMBER GENERATED ALWAYS AS IDENTITY NOT NULL,\n"
    "    STG_RUN_ID        NUMBER         NOT NULL,\n"
    "    STG_STATUS        VARCHAR2(20)   DEFAULT 'PENDING' NOT NULL,\n"
    "    STG_ACTION        VARCHAR2(20),\n"
    "    STG_REJECT_REASON VARCHAR2(500),\n"
    "    STG_RECORD_HASH   VARCHAR2(64),\n"
    "    MY_ENTITY_ID      VARCHAR2(50),\n"
    "    ATTRIBUTE_1       VARCHAR2(100),\n"
    "    ATTRIBUTE_2       DATE,\n"
    "    REPORTING_DATE    DATE,\n"
    "    CONSTRAINT PK_STG_ME PRIMARY KEY (STG_ID)\n"
    ");\n"
    "CREATE INDEX GPC_DM.IDX_STG_ME_RUN ON GPC_DM.STG_MY_ENTITY(STG_RUN_ID, STG_STATUS);"
)

h2('Step 2 — Transform Package (required signature)')
add_code(
    "CREATE OR REPLACE PACKAGE GPC_DM.PKG_ETL_TRANSFORM_MY_ENTITY AS\n"
    "    PROCEDURE transform(\n"
    "        p_run_id         IN NUMBER,\n"
    "        p_entity_id      IN NUMBER,\n"
    "        p_from_watermark IN DATE,\n"
    "        p_to_watermark   IN DATE\n"
    "    );\n"
    "END PKG_ETL_TRANSFORM_MY_ENTITY;\n/\n\n"
    "-- Key requirements in the body:\n"
    "-- 1. DELETE FROM STG_MY_ENTITY WHERE STG_STATUS = 'PENDING';  (idempotency)\n"
    "-- 2. INSERT INTO STG_MY_ENTITY ... SELECT ... WHERE REPORTING_DATE > p_from_watermark ...\n"
    "-- 3. Compute STANDARD_HASH over tracked columns for STG_RECORD_HASH\n"
    "-- 4. Call PKG_ETL_LOGGER.log_step / end_step for observability"
)

h2('Step 3 — Register Metadata')
add_code(
    "DECLARE\n"
    "    v_ss_id  NUMBER;  v_eid  NUMBER;  v_mid  NUMBER;\n"
    "BEGIN\n"
    "    SELECT SS_ID INTO v_ss_id FROM GPC_DM.ETL_SOURCE_SYSTEM WHERE SS_NAME = 'KBR_IHUB';\n\n"
    "    INSERT INTO GPC_DM.ETL_ENTITY (SS_ID, ENTITY_NAME, SOURCE_TABLE, WATERMARK_COLUMN, WATERMARK_TYPE)\n"
    "    VALUES (v_ss_id, 'MY_ENTITY', 'MY_SCHEMA.MY_TABLE', 'REPORTING_DATE', 'DATE')\n"
    "    RETURNING ENTITY_ID INTO v_eid;\n\n"
    "    INSERT INTO GPC_DM.ETL_CONTROL (ENTITY_ID, STATUS) VALUES (v_eid, 'IDLE');\n\n"
    "    INSERT INTO GPC_DM.ETL_TARGET_MAPPING (\n"
    "        ENTITY_ID, TARGET_TABLE, STAGING_TABLE, LOAD_TYPE, SURROGATE_KEY_COL, SURROGATE_SEQ_NAME\n"
    "    ) VALUES (\n"
    "        v_eid, 'GPC_DM.DIM_MY_ENTITY', 'GPC_DM.STG_MY_ENTITY', 'SCD2',\n"
    "        'DIM_ME_ID', 'GPC_DM.SEQ_DIM_MY_ENTITY'\n"
    "    ) RETURNING MAPPING_ID INTO v_mid;\n\n"
    "    -- Business key column (IS_BUSINESS_KEY = Y)\n"
    "    INSERT INTO GPC_DM.ETL_COLUMN_MAPPING\n"
    "        (MAPPING_ID, SOURCE_COLUMN, TARGET_COLUMN, IS_BUSINESS_KEY, IS_TRACKED, COLUMN_ORDER)\n"
    "    VALUES (v_mid, 'MY_ENTITY_ID', 'MY_ENTITY_ID', 'Y', 'N', 10);\n\n"
    "    -- Tracked attribute (IS_TRACKED = Y → included in change detection hash)\n"
    "    INSERT INTO GPC_DM.ETL_COLUMN_MAPPING\n"
    "        (MAPPING_ID, SOURCE_COLUMN, TARGET_COLUMN, IS_BUSINESS_KEY, IS_TRACKED, COLUMN_ORDER)\n"
    "    VALUES (v_mid, 'ATTRIBUTE_1', 'ATTRIBUTE_1', 'N', 'Y', 20);\n\n"
    "    -- Validation rule\n"
    "    INSERT INTO GPC_DM.ETL_VALIDATION_RULE\n"
    "        (MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, ERROR_ACTION)\n"
    "    VALUES (v_mid, 'MY_ENTITY_ID_REQUIRED', 'NOT_NULL', 'MY_ENTITY_ID', 'REJECT');\n\n"
    "    COMMIT;\n"
    "END;\n/"
)

h2('Step 4 — Add Dispatch Branch in Orchestrator')
para('Edit packages/15_pkg_etl_orchestrator.sql — add a WHEN branch inside dispatch_transform:')
add_code(
    "WHEN 'MY_ENTITY' THEN\n"
    "    GPC_DM.PKG_ETL_TRANSFORM_MY_ENTITY.transform(\n"
    "        p_run_id, p_entity_id, p_from_watermark, p_to_watermark);\n\n"
    "-- Then recompile:\n"
    "@packages/15_pkg_etl_orchestrator.sql"
)

h2('Step 5 — Test')
add_code(
    "EXEC GPC_DM.PKG_ETL_ORCHESTRATOR.run_entity('MY_ENTITY');\n\n"
    "SELECT * FROM GPC_DM.V_ETL_RUN_SUMMARY WHERE ROWNUM <= 5;\n"
    "SELECT COUNT(*) FROM GPC_DM.DIM_MY_ENTITY WHERE IS_CURRENT = 'Y';\n"
    "SELECT COUNT(*) FROM GPC_DM.STG_MY_ENTITY  WHERE STG_STATUS = 'REJECTED';"
)


# ════════════════════════════════════════════════════════════════════════════
# 12. MODIFY EXISTING ENTITY
# ════════════════════════════════════════════════════════════════════════════
h1('12. How to Modify an Existing Entity')

h2('Add a New Column')
add_code(
    "-- 1. Alter tables\n"
    "ALTER TABLE GPC_DM.DIM_STAFFING_SCHEDULE ADD (NEW_COL VARCHAR2(100));\n"
    "ALTER TABLE GPC_DM.STG_STAFFING_SCHEDULE ADD (NEW_COL VARCHAR2(100));\n\n"
    "-- 2. Register column mapping\n"
    "INSERT INTO GPC_DM.ETL_COLUMN_MAPPING\n"
    "    (MAPPING_ID, SOURCE_COLUMN, TARGET_COLUMN, IS_BUSINESS_KEY, IS_TRACKED, COLUMN_ORDER)\n"
    "VALUES (<mapping_id>, 'SRC_COL', 'NEW_COL', 'N', 'Y', <next_order>);\n"
    "COMMIT;\n\n"
    "-- 3. Update the transform package INSERT SELECT to include the new column\n"
    "-- 4. If IS_TRACKED = Y: add column to STANDARD_HASH in the transform package"
)

h2('Remove a Column (Safe)')
add_code(
    "-- Mark inactive (preserves audit history)\n"
    "UPDATE GPC_DM.ETL_COLUMN_MAPPING SET IS_ACTIVE = 'N'\n"
    "WHERE  TARGET_COLUMN = 'OLD_COL' AND MAPPING_ID = <mapping_id>;\n"
    "COMMIT;\n\n"
    "-- Remove from transform package INSERT SELECT\n"
    "-- Optionally: ALTER TABLE ... DROP COLUMN OLD_COL;"
)

h2('Change Load Type')
add_code(
    "UPDATE GPC_DM.ETL_TARGET_MAPPING SET LOAD_TYPE = 'INCREMENTAL'\n"
    "WHERE  TARGET_TABLE = 'GPC_DM.DIM_MY_ENTITY';\n"
    "COMMIT;\n"
    "-- No package recompilation required"
)


# ════════════════════════════════════════════════════════════════════════════
# 13. ADD / MODIFY VALIDATION RULES
# ════════════════════════════════════════════════════════════════════════════
h1('13. How to Add or Modify Validation Rules')

para('Validation rules are purely metadata — no package recompilation needed to add or change them.')

h2('Add a NOT_NULL Rule')
add_code(
    "INSERT INTO GPC_DM.ETL_VALIDATION_RULE\n"
    "    (MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, ERROR_ACTION)\n"
    "VALUES (<mapping_id>, 'MY_COL_REQUIRED', 'NOT_NULL', 'MY_COLUMN', 'REJECT');\n"
    "COMMIT;"
)

h2('Add a CHECK Rule (enum / condition)')
add_code(
    "INSERT INTO GPC_DM.ETL_VALIDATION_RULE\n"
    "    (MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, DERIVED_SQL, ERROR_ACTION)\n"
    "VALUES (\n"
    "    <mapping_id>,\n"
    "    'MY_COL_VALUE_CHECK',\n"
    "    'CHECK',\n"
    "    'MY_COLUMN',\n"
    "    'MY_COLUMN IS NULL OR MY_COLUMN IN (''VALUE_A'',''VALUE_B'')',\n"
    "    'REJECT'\n"
    ");\n"
    "COMMIT;\n\n"
    "-- DERIVED_SQL must be a valid Oracle WHERE-clause predicate.\n"
    "-- The validator executes: WHERE NOT (<DERIVED_SQL>)"
)

h2('Add a DERIVED Rule (compute value, then validate not NULL)')
add_code(
    "INSERT INTO GPC_DM.ETL_VALIDATION_RULE\n"
    "    (MAPPING_ID, RULE_NAME, RULE_TYPE, COLUMN_NAME, DERIVED_SQL, ERROR_ACTION)\n"
    "VALUES (\n"
    "    <mapping_id>,\n"
    "    'DT_PERIOD_DERIVED',\n"
    "    'DERIVED',\n"
    "    'DT_PERIOD',\n"
    "    'TO_CHAR(PERIOD_START_DATE, ''YYYYMM'')',\n"
    "    'REJECT'\n"
    ");\n"
    "COMMIT;"
)

h2('Rule Type Summary')
add_table(
    ['RULE_TYPE', 'DERIVED_SQL content', 'Behaviour'],
    [
        ['NOT_NULL', '(leave NULL)', 'Rejects rows where COLUMN_NAME IS NULL'],
        ['DERIVED',  'SQL expression (e.g. TO_CHAR(DATE_COL,\'YYYYMM\'))', 'Populates NULL column; then rejects if still NULL'],
        ['CHECK',    'WHERE predicate (e.g. COL IN (\'A\',\'B\'))', 'Rejects rows where NOT (predicate) is true'],
    ],
    col_widths=[2.5, 6.5, 7.0]
)

h2('Escalate to FAIL (abort on any violation)')
add_code(
    "UPDATE GPC_DM.ETL_VALIDATION_RULE SET ERROR_ACTION = 'FAIL'\n"
    "WHERE  RULE_NAME = 'MY_COL_REQUIRED';\n"
    "COMMIT;"
)

h2('Execution Order')
para('Rules execute in RULE_ID ascending order (assigned by sequence at insert time). Recommended order: DERIVED rules first, then NOT_NULL, then CHECK — so derivation runs before value checks.')


# ════════════════════════════════════════════════════════════════════════════
# 14. VALIDATION RULES REFERENCE
# ════════════════════════════════════════════════════════════════════════════
h1('14. Validation Rules Reference')

rule_sections = [
    ('STAFFING_SCHEDULE → DIM_STAFFING_SCHEDULE', [
        ('PROJECT_ID_REQUIRED',  'NOT_NULL', 'PROJECT_ID',         'Must not be NULL',                                            'REJECT'),
        ('POSITION_ID_REQUIRED', 'NOT_NULL', 'POSITION_ID',        'Must not be NULL',                                            'REJECT'),
        ('SCHEDULE_TYPE_CHECK',  'CHECK',    'SCHEDULE_TYPE',      "NULL or IN ('Exempt','Non-Exempt','Exempt Agency','Non-Exempt Agency')", 'REJECT'),
        ('SCHEDULE_STATUS_CHECK','CHECK',    'SCHEDULE_STATUS',    "NULL or IN ('Open','Filled','Canceled','Approved','Rejected','Withdrawn','Pending')", 'REJECT'),
        ('SCHEDULE_DATE_ORDER',  'CHECK',    'SCHEDULE_START_DATE','NULL or START_DATE <= END_DATE',                               'REJECT'),
    ]),
    ('STAFFING_SCHEDULE → DIM_STAFFING_TIMELINE', [
        ('ALLOCATED_HOURS_REQUIRED',     'NOT_NULL','ALLOCATED_HOURS',   'Must not be NULL',                              'REJECT'),
        ('ALLOCATED_HOURS_NON_NEGATIVE', 'CHECK',   'ALLOCATED_HOURS',   'NULL or >= 0',                                  'REJECT'),
        ('PERIOD_DATE_ORDER',            'CHECK',   'PERIOD_START_DATE', 'NULL or PERIOD_START_DATE <= PERIOD_END_DATE',  'REJECT'),
    ]),
    ('COST → DIM_COST', [
        ('PROJECT_ID_REQUIRED', 'NOT_NULL', 'PROJECT_ID',    'Must not be NULL',                                      'REJECT'),
        ('COST_TYPE_REQUIRED',  'NOT_NULL', 'COST_TYPE',     'Must not be NULL',                                      'REJECT'),
        ('COST_CATEGORY_CHECK', 'CHECK',    'COST_CATEGORY', "NULL or UPPER(COST_CATEGORY) IN ('COMPANY','CLIENT')",  'REJECT'),
    ]),
    ('COST → DIM_TIMELINE_COST', [
        ('AMOUNT_REQUIRED',        'NOT_NULL','AMOUNT',        'Must not be NULL',                              'REJECT'),
        ('FORECAST_TYPE_REQUIRED', 'NOT_NULL','FORECAST_TYPE', 'Must not be NULL',                              'REJECT'),
        ('CURVE_TYPE_VALUE_CHECK', 'CHECK',   'CURVE_TYPE',    "NULL or IN ('ACTUAL','BUDGET','FORECAST')",     'REJECT'),
    ]),
]

for section_title, rules in rule_sections:
    h2(section_title)
    add_table(
        ['Rule Name', 'Type', 'Column', 'Condition', 'Action'],
        [list(r) for r in rules],
        col_widths=[4.5, 2.0, 3.5, 5.5, 1.5]
    )


# ════════════════════════════════════════════════════════════════════════════
# 15. METADATA TABLE REFERENCE
# ════════════════════════════════════════════════════════════════════════════
h1('15. Metadata Table Reference')

tables = [
    ('ETL_SOURCE_SYSTEM', [
        ('SS_ID', 'NUMBER', 'Surrogate PK'),
        ('SS_NAME', 'VARCHAR2(100)', 'Unique name, e.g. KBR_IHUB'),
        ('SS_SCHEMA', 'VARCHAR2(100)', 'Oracle schema prefix for source tables'),
        ('IS_ACTIVE', 'VARCHAR2(1)', 'Y/N'),
    ]),
    ('ETL_ENTITY', [
        ('ENTITY_ID', 'NUMBER', 'Surrogate PK'),
        ('SS_ID', 'NUMBER', 'FK to ETL_SOURCE_SYSTEM'),
        ('ENTITY_NAME', 'VARCHAR2(100)', 'Unique; used in run_entity() calls'),
        ('SOURCE_TABLE', 'VARCHAR2(200)', 'Fully qualified source table name'),
        ('WATERMARK_COLUMN', 'VARCHAR2(100)', 'Column used for incremental detection'),
        ('WATERMARK_TYPE', 'VARCHAR2(20)', 'DATE / TIMESTAMP / NUMBER'),
    ]),
    ('ETL_TARGET_MAPPING', [
        ('MAPPING_ID', 'NUMBER', 'Surrogate PK'),
        ('ENTITY_ID', 'NUMBER', 'FK to ETL_ENTITY'),
        ('TARGET_TABLE', 'VARCHAR2(200)', 'Fully qualified target DIM table'),
        ('STAGING_TABLE', 'VARCHAR2(200)', 'Fully qualified staging STG table'),
        ('LOAD_TYPE', 'VARCHAR2(20)', 'SCD1 / SCD2 / INCREMENTAL / APPEND'),
        ('SURROGATE_KEY_COL', 'VARCHAR2(100)', 'Target column for surrogate key'),
        ('SURROGATE_SEQ_NAME', 'VARCHAR2(200)', 'Sequence to call for surrogate key values'),
        ('LOAD_ORDER', 'NUMBER', 'Execution order within entity (lower = first)'),
    ]),
    ('ETL_COLUMN_MAPPING', [
        ('CM_ID', 'NUMBER', 'Surrogate PK'),
        ('MAPPING_ID', 'NUMBER', 'FK to ETL_TARGET_MAPPING'),
        ('SOURCE_COLUMN', 'VARCHAR2(100)', 'Column name in staging table'),
        ('TARGET_COLUMN', 'VARCHAR2(100)', 'Column name in target DIM table'),
        ('IS_BUSINESS_KEY', 'VARCHAR2(1)', 'Y = used in SCD2/MERGE join condition'),
        ('IS_TRACKED', 'VARCHAR2(1)', 'Y = included in RECORD_HASH change detection'),
        ('COLUMN_ORDER', 'NUMBER', 'Hash concatenation order — must be consistent'),
    ]),
    ('ETL_VALIDATION_RULE', [
        ('RULE_ID', 'NUMBER', 'Surrogate PK; determines execution order (ascending)'),
        ('MAPPING_ID', 'NUMBER', 'FK to ETL_TARGET_MAPPING'),
        ('RULE_NAME', 'VARCHAR2(100)', 'Descriptive name; appears in STG_REJECT_REASON'),
        ('RULE_TYPE', 'VARCHAR2(20)', 'NOT_NULL / DERIVED / CHECK / CUSTOM'),
        ('COLUMN_NAME', 'VARCHAR2(100)', 'Staging column the rule operates on'),
        ('DERIVED_SQL', 'VARCHAR2(4000)', 'SQL expression (DERIVED) or WHERE predicate (CHECK)'),
        ('ERROR_ACTION', 'VARCHAR2(10)', 'REJECT = continue; FAIL = abort run'),
        ('IS_ACTIVE', 'VARCHAR2(1)', 'Y/N; disable without deleting'),
    ]),
    ('ETL_CONTROL', [
        ('ENTITY_ID', 'NUMBER', 'FK to ETL_ENTITY (unique)'),
        ('LAST_WATERMARK', 'DATE', 'Watermark after last successful run; NULL = never run (full load)'),
        ('STATUS', 'VARCHAR2(20)', 'IDLE / RUNNING / FAILED / DISABLED'),
        ('LAST_RUN_DATE', 'DATE', 'Timestamp of last run attempt'),
        ('NOTES', 'VARCHAR2(500)', 'Operator notes'),
    ]),
    ('ETL_RUN_LOG', [
        ('RUN_ID', 'NUMBER', 'Unique run identifier'),
        ('ENTITY_ID', 'NUMBER', 'Entity being processed'),
        ('MAPPING_ID', 'NUMBER', 'NULL = entity-level entry; populated = mapping-level'),
        ('STATUS', 'VARCHAR2(20)', 'RUNNING / SUCCESS / FAILED / PARTIAL'),
        ('ROWS_READ', 'NUMBER', 'Rows extracted from source'),
        ('ROWS_INSERTED', 'NUMBER', 'New rows written to target'),
        ('ROWS_UPDATED', 'NUMBER', 'Rows expired (SCD2) or updated'),
        ('ROWS_REJECTED', 'NUMBER', 'Rows that failed validation'),
    ]),
    ('ETL_ERROR_LOG', [
        ('ERR_ID', 'NUMBER', 'Surrogate PK'),
        ('RUN_ID', 'NUMBER', 'FK to ETL_RUN_LOG'),
        ('ERROR_CODE', 'VARCHAR2(50)', 'VAL_NOT_NULL / VAL_DERIVED_NULL / VAL_CHECK_FAILED / SCD2_POST_LOAD_DUP'),
        ('ERROR_MESSAGE', 'VARCHAR2(4000)', 'Full error description'),
        ('RECORD_KEY', 'VARCHAR2(500)', 'Serialised business key of the failing row'),
        ('RECORD_DATA', 'CLOB', 'Full staging row for post-mortem investigation'),
    ]),
]

for tbl_name, cols in tables:
    h2(tbl_name)
    add_table(['Column', 'Type', 'Description'], cols, col_widths=[4.0, 3.0, 9.0])


# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(r'c:\Aralytiks\AR-Dev\EDWH_Framework\SOLUTION_MANUAL.docx')
print("Done: SOLUTION_MANUAL.docx created.")
